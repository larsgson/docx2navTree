#!/usr/bin/env python3
"""
Reference Script: Split Book to Markdown Chapters
==================================================

This reference script demonstrates how the document conversion system works by
generating Markdown (.md) files instead of JSON. Each chapter is saved as a
separate .md file with formatting that closely matches the chapter-viewer display.

Features:
- Converts chapters to readable Markdown format
- Preserves text formatting (bold, italic, font sizes)
- Includes tables in Markdown table format
- Embeds images with proper references
- Handles footnotes
- Creates chapter navigation structure

This is a reference implementation based on build_book.py logic.

Usage:
    python3 split_to_md_chapters.py

Output:
    Creates markdown_chapters/ directory with:
    - chapter_01.md, chapter_02.md, etc.
    - pictures/ subdirectories with images
    - index.md with navigation
"""

import json
import os
import re
import shutil
import subprocess
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

# ============================================================================
# Configuration
# ============================================================================

INPUT_DOCX = "English HAH Word Apr 6 2024.docx"
MARKDOWN_DIR = "markdown_chapters"

# ============================================================================
# Color Output for Terminal
# ============================================================================


class Colors:
    HEADER = "\033[95m"
    BLUE = "\033[94m"
    CYAN = "\033[96m"
    GREEN = "\033[92m"
    YELLOW = "\033[93m"
    RED = "\033[91m"
    END = "\033[0m"
    BOLD = "\033[1m"


def print_header(text):
    print(f"\n{Colors.BOLD}{Colors.HEADER}{'=' * 70}{Colors.END}")
    print(f"{Colors.BOLD}{Colors.HEADER}{text}{Colors.END}")
    print(f"{Colors.BOLD}{Colors.HEADER}{'=' * 70}{Colors.END}\n")


def print_step(step_num, text):
    print(f"\n{Colors.BOLD}{Colors.BLUE}Step {step_num}: {text}{Colors.END}")
    print(f"{Colors.BLUE}{'-' * 70}{Colors.END}")


def print_success(text):
    print(f"{Colors.GREEN}✅ {text}{Colors.END}")


def print_warning(text):
    print(f"{Colors.YELLOW}⚠️  {text}{Colors.END}")


def print_error(text):
    print(f"{Colors.RED}❌ {text}{Colors.END}")


def print_info(text):
    print(f"{Colors.CYAN}ℹ️  {text}{Colors.END}")


# ============================================================================
# Image Conversion Functions (from build_book.py)
# ============================================================================


def check_imagemagick():
    """Check if ImageMagick is installed."""
    try:
        result = subprocess.run(
            ["convert", "-version"], capture_output=True, text=True, timeout=5
        )
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


def check_ghostscript():
    """Check if Ghostscript is installed."""
    try:
        result = subprocess.run(
            ["gs", "--version"], capture_output=True, text=True, timeout=5
        )
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


def check_file_type(filepath):
    """Check if a file is a Windows Metafile."""
    try:
        result = subprocess.run(
            ["file", "-b", str(filepath)], capture_output=True, text=True, timeout=5
        )
        output = result.stdout.strip().lower()
        return "windows metafile" in output or "wmf" in output
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


def convert_wmf_to_png(filepath):
    """Convert a WMF file to PNG using ImageMagick.

    ImageMagick will automatically handle the conversion chain:
    WMF -> (LibreOffice) -> PDF -> (Ghostscript) -> PNG
    """
    # Use .wmf extension (not .wmf.tmp) so ImageMagick recognizes the format
    temp_wmf = filepath.parent / f"{filepath.stem}_temp.wmf"
    temp_png = filepath.parent / f"{filepath.stem}_temp.png"

    try:
        # Rename file to .wmf extension so ImageMagick recognizes the format
        shutil.copy2(filepath, temp_wmf)

        # Convert WMF to PNG - ImageMagick handles PDF intermediate automatically
        result = subprocess.run(
            ["magick", str(temp_wmf), str(temp_png)],
            capture_output=True,
            text=True,
            timeout=30,
        )

        if result.returncode == 0 and temp_png.exists():
            # Verify it's actually a PNG
            check_result = subprocess.run(
                ["file", "-b", str(temp_png)],
                capture_output=True,
                text=True,
            )
            if "PNG" not in check_result.stdout:
                # Not a proper PNG, fail
                if temp_wmf.exists():
                    temp_wmf.unlink()
                if temp_png.exists():
                    temp_png.unlink()
                return False

            # Backup original WMF file
            backup_path = filepath.with_suffix(".wmf.backup")
            shutil.copy2(filepath, backup_path)

            # Replace with converted PNG
            shutil.move(str(temp_png), str(filepath))

            # Clean up temp WMF file
            if temp_wmf.exists():
                temp_wmf.unlink()

            return True
        else:
            # Clean up temp files
            if temp_wmf.exists():
                temp_wmf.unlink()
            if temp_png.exists():
                temp_png.unlink()
            return False

    except Exception as e:
        # Clean up temp files
        if temp_wmf.exists():
            temp_wmf.unlink()
        if temp_png.exists():
            temp_png.unlink()
        return False


# ============================================================================
# Utility Functions (from build_book.py)
# ============================================================================


def clean_toc_formatting(text):
    """Remove TOC formatting from title text (dots and page numbers)."""
    if not text:
        return text

    # Remove trailing dots and page numbers
    cleaned = re.sub(r"\s*\.+\s*\.+.*?\d+\s*$", "", text)
    cleaned = re.sub(r"\s+\.{2,}.*$", "", cleaned)

    return cleaned.strip()


def is_toc_entry(text):
    """Check if text looks like a TOC entry (has multiple dots)."""
    return text and "....." in text


def should_skip_toc_paragraph(paragraph):
    """Determine if a paragraph should be skipped as TOC content."""
    text = paragraph.text.strip()
    return is_toc_entry(text)


def is_chapter_heading(paragraph):
    """Check if paragraph is a chapter heading (N.0 format)."""
    text = paragraph.text.strip()

    # Match N.0 pattern
    if not re.match(r"^\d+\.0\s", text):
        return False

    # Check for bold formatting
    for run in paragraph.runs:
        if run.bold and run.text.strip():
            return True

    return False


def extract_chapter_number(paragraph):
    """Extract chapter number from heading."""
    text = paragraph.text.strip()
    match = re.match(r"^(\d+)\.0\s", text)
    if match:
        return int(match.group(1))
    return None


def extract_section_number(paragraph):
    """Extract section number (e.g., 1.1, 1.2)."""
    text = paragraph.text.strip()
    match = re.match(r"^(\d+)\.(\d+)", text)
    return match.group(0) if match else None


def is_first_section_heading(paragraph):
    """Check if paragraph is first section in chapter (N.1 format)."""
    text = paragraph.text.strip()

    if not re.match(r"^\d+\.1\s", text):
        return False

    for run in paragraph.runs:
        if run.bold and run.text.strip():
            return True

    return False


def get_footnote_references_from_paragraph(paragraph):
    """Get all footnote reference IDs from a paragraph."""
    footnote_refs = []
    try:
        refs = paragraph._element.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnoteReference"
        )
        for footnote_ref in refs:
            footnote_id = footnote_ref.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id"
            )
            if footnote_id:
                footnote_refs.append(footnote_id)
    except:
        pass
    return footnote_refs


def get_document_elements(doc):
    """Get all document elements (paragraphs and tables) in order."""
    elements = []
    for element in doc.element.body:
        if isinstance(element, CT_P):
            para = Paragraph(element, doc)
            elements.append(("paragraph", para))
        elif isinstance(element, CT_Tbl):
            table = Table(element, doc)
            elements.append(("table", table))
    return elements


# ============================================================================
# Markdown Generation Functions
# ============================================================================


def escape_markdown(text):
    """Escape special Markdown characters in text."""
    # Don't escape in code blocks or if already escaped
    if not text:
        return text

    # Escape pipe characters for tables
    text = text.replace("|", "\\|")

    return text


def format_run_markdown(run):
    """Format a text run with Markdown syntax."""
    text = run.text

    if not text:
        return ""

    # Apply formatting
    if run.bold and run.italic:
        text = f"***{text}***"
    elif run.bold:
        text = f"**{text}**"
    elif run.italic:
        text = f"*{text}*"

    if run.underline:
        # Markdown doesn't have native underline, use HTML
        text = f"<u>{text}</u>"

    return text


def extract_paragraph_markdown(
    paragraph, image_counter, image_dir, source_doc, chapter_num
):
    """Extract paragraph and convert to Markdown format."""

    markdown_lines = []
    para_text = paragraph.text.strip()

    # Skip empty paragraphs
    if not para_text and not has_images(paragraph):
        return "", image_counter

    # Check if it's a heading
    section_num = extract_section_number(paragraph)
    if section_num:
        # Check heading level based on formatting
        is_bold = any(run.bold for run in paragraph.runs if run.text.strip())

        if is_bold:
            if section_num.endswith(".0"):
                # Chapter heading (H1)
                markdown_lines.append(f"# {para_text}\n")
            elif section_num.count(".") == 1:
                # Section heading (H2)
                markdown_lines.append(f"## {para_text}\n")
            elif section_num.count(".") == 2:
                # Subsection heading (H3)
                markdown_lines.append(f"### {para_text}\n")
            else:
                # Sub-subsection heading (H4)
                markdown_lines.append(f"#### {para_text}\n")
        else:
            # Not a heading, just bold section number
            formatted_text = ""
            for run in paragraph.runs:
                formatted_text += format_run_markdown(run)
            markdown_lines.append(f"{formatted_text}\n")
    else:
        # Regular paragraph with formatting
        formatted_text = ""
        for run in paragraph.runs:
            formatted_text += format_run_markdown(run)

        if formatted_text.strip():
            markdown_lines.append(f"{formatted_text}\n")

    # Extract inline images
    for run in paragraph.runs:
        inline_shapes = run._element.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
        )

        for shape in inline_shapes:
            blip = shape.find(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )
            if blip is not None:
                image_rId = blip.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )

                if image_rId:
                    try:
                        image_part = source_doc.part.related_parts[image_rId]
                        image_bytes = image_part.blob

                        # Determine image format
                        content_type = image_part.content_type
                        if "png" in content_type:
                            ext = "png"
                        elif "jpeg" in content_type or "jpg" in content_type:
                            ext = "jpg"
                        elif "gif" in content_type:
                            ext = "gif"
                        elif "wmf" in content_type or "x-wmf" in content_type:
                            ext = "wmf"
                        else:
                            ext = "png"

                        # Save image
                        image_filename = f"image_{image_counter:04d}.{ext}"
                        image_path = os.path.join(image_dir, image_filename)

                        os.makedirs(image_dir, exist_ok=True)
                        with open(image_path, "wb") as f:
                            f.write(image_bytes)

                        # Convert WMF to PNG if needed
                        if ext == "wmf":
                            image_path_obj = Path(image_path)
                            print_info(f"    Converting WMF image {image_counter}...")
                            if convert_wmf_to_png(image_path_obj):
                                # Rename the converted file from .wmf to .png
                                png_filename = f"image_{image_counter:04d}.png"
                                png_path = os.path.join(image_dir, png_filename)

                                # The convert_wmf_to_png replaces the .wmf file with PNG data
                                # but keeps the .wmf extension, so we need to rename it
                                if os.path.exists(image_path):
                                    os.rename(image_path, png_path)

                                # Update filename to PNG
                                image_filename = png_filename
                                image_relative_path = f"pictures/{image_filename}"
                                print_success(
                                    f"      Converted to PNG: {image_filename}"
                                )
                            else:
                                # Conversion failed, keep as WMF
                                image_relative_path = f"pictures/{image_filename}"
                                print_warning(
                                    f"      WMF conversion failed for {image_filename}"
                                )
                        else:
                            image_relative_path = f"pictures/{image_filename}"

                        # Add image reference in Markdown
                        markdown_lines.append(
                            f"\n![Image {image_counter}]({image_relative_path})\n"
                        )

                        image_counter += 1

                    except Exception as e:
                        print_warning(f"Could not extract image: {e}")
                        markdown_lines.append(
                            f"\n*[Image {image_counter} - extraction failed]*\n"
                        )

    # Add footnote references if any
    footnote_refs = get_footnote_references_from_paragraph(paragraph)
    if footnote_refs:
        footnote_text = ", ".join([f"[^{ref}]" for ref in footnote_refs])
        markdown_lines.append(f" {footnote_text}")

    return "\n".join(markdown_lines), image_counter


def has_images(paragraph):
    """Check if paragraph contains images."""
    for run in paragraph.runs:
        inline_shapes = run._element.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
        )
        if inline_shapes:
            return True
    return False


def extract_table_markdown(table):
    """Extract table and convert to Markdown table format."""

    if not table.rows:
        return ""

    markdown_lines = []

    # Determine number of columns
    num_cols = len(table.columns) if table.columns else len(table.rows[0].cells)

    # Process table rows
    for row_idx, row in enumerate(table.rows):
        row_cells = []

        for cell in row.cells:
            # Extract cell text with formatting
            cell_text = ""
            for para in cell.paragraphs:
                for run in para.runs:
                    cell_text += format_run_markdown(run)

            # Clean up and escape
            cell_text = cell_text.strip()
            cell_text = escape_markdown(cell_text) if cell_text else " "
            row_cells.append(cell_text)

        # Add row to markdown
        markdown_lines.append("| " + " | ".join(row_cells) + " |")

        # Add separator after first row (header)
        if row_idx == 0:
            separator = "| " + " | ".join(["---"] * len(row_cells)) + " |"
            markdown_lines.append(separator)

    return "\n" + "\n".join(markdown_lines) + "\n"


def split_into_sections(chapter_elements, chapter_num, chapter_title):
    """Split chapter elements into sections based on N.X headings."""

    sections = []
    current_section = {
        "number": f"{chapter_num}.0",
        "title": chapter_title,
        "elements": [],
    }

    for elem_type, elem in chapter_elements:
        if elem_type == "paragraph":
            section_num = extract_section_number(elem)

            # Check if this is a new section (N.X where X > 0)
            if section_num and section_num != f"{chapter_num}.0":
                # Save previous section if it has content
                if current_section["elements"]:
                    sections.append(current_section)

                # Start new section
                current_section = {
                    "number": section_num,
                    "title": elem.text.strip(),
                    "elements": [(elem_type, elem)],
                }
            else:
                current_section["elements"].append((elem_type, elem))
        else:
            current_section["elements"].append((elem_type, elem))

    # Add final section
    if current_section["elements"]:
        sections.append(current_section)

    return sections


# ============================================================================
# Main Conversion Function
# ============================================================================


def convert_book_to_markdown(input_file, output_dir):
    """Convert Word document to Markdown files, one per chapter."""

    print_header("MARKDOWN CHAPTER GENERATOR")
    print_info(f"Input: {input_file}")
    print_info(f"Output: {output_dir}/")

    # Check for ImageMagick (for WMF conversion)
    has_imagemagick = check_imagemagick()
    has_ghostscript = check_ghostscript()

    if has_imagemagick and has_ghostscript:
        print_success("ImageMagick and Ghostscript detected - WMF conversion enabled")
    elif has_imagemagick:
        print_warning(
            "ImageMagick found but Ghostscript missing - WMF conversion may fail"
        )
    else:
        print_warning("ImageMagick not found - WMF images will not be converted")

    # Create output directory
    os.makedirs(output_dir, exist_ok=True)

    # Load document
    print_step(1, "Loading Word document")
    doc = Document(input_file)
    print_success(f"Loaded document with {len(doc.paragraphs)} paragraphs")

    # Get all elements
    print_step(2, "Extracting document structure")
    elements = get_document_elements(doc)
    print_success(f"Found {len(elements)} total elements")

    # Identify chapters
    print_step(3, "Identifying chapters")

    chapters = []
    current_chapter = None
    current_chapter_title = None
    current_chapter_elements = []
    appendix_start = False

    for elem_type, elem in elements:
        if elem_type == "paragraph":
            # Skip TOC entries
            if should_skip_toc_paragraph(elem):
                continue

            # Check for chapter heading (N.0)
            if is_chapter_heading(elem):
                chapter_num = extract_chapter_number(elem)

                # Save previous chapter
                if current_chapter is not None and current_chapter_elements:
                    chapters.append(
                        {
                            "number": current_chapter,
                            "title": current_chapter_title,
                            "elements": current_chapter_elements,
                        }
                    )
                    print_info(
                        f"  Found Chapter {current_chapter}: {current_chapter_title[:50]}..."
                    )

                # Start new chapter
                current_chapter = chapter_num
                current_chapter_title = elem.text.strip()
                current_chapter_elements = [(elem_type, elem)]

            # Check for first section (N.1) - handles appendix-style chapters
            elif is_first_section_heading(elem) and current_chapter is None:
                section_text = elem.text.strip()
                match = re.match(r"^(\d+)\.1\s", section_text)
                if match:
                    chapter_num = int(match.group(1))

                    # Start new appendix-style chapter
                    current_chapter = chapter_num
                    current_chapter_title = f"{chapter_num}.0 {section_text.split(' ', 1)[1] if ' ' in section_text else 'Chapter'}"
                    current_chapter_elements = [(elem_type, elem)]
                    appendix_start = True
            else:
                if current_chapter is not None:
                    current_chapter_elements.append((elem_type, elem))
        else:
            # Table or other element
            if current_chapter is not None:
                current_chapter_elements.append((elem_type, elem))

    # Save last chapter
    if current_chapter is not None and current_chapter_elements:
        chapters.append(
            {
                "number": current_chapter,
                "title": current_chapter_title,
                "elements": current_chapter_elements,
            }
        )
        print_info(
            f"  Found Chapter {current_chapter}: {current_chapter_title[:50]}..."
        )

    print_success(f"Identified {len(chapters)} chapters")

    # Generate Markdown files
    print_step(4, "Generating Markdown files")

    chapter_index = []
    total_sections = 0
    total_paragraphs = 0
    total_tables = 0
    total_images = 0
    total_wmf_converted = 0

    for chapter in chapters:
        chapter_num = chapter["number"]
        chapter_title = chapter["title"]
        chapter_elements = chapter["elements"]

        # Create chapter directory
        chapter_dir = os.path.join(output_dir, f"chapter_{chapter_num:02d}")
        os.makedirs(chapter_dir, exist_ok=True)

        # Create pictures directory
        pictures_dir = os.path.join(chapter_dir, "pictures")
        os.makedirs(pictures_dir, exist_ok=True)

        # Split into sections
        sections = split_into_sections(chapter_elements, chapter_num, chapter_title)

        print_info(f"  Chapter {chapter_num}: {len(sections)} sections")

        # Generate markdown for each section
        section_files = []
        chapter_image_counter = 1

        for section in sections:
            section_num = section["number"]
            section_title = section["title"]
            section_elements = section["elements"]

            # Create markdown content
            markdown_content = []
            markdown_content.append(f"# {section_title}\n\n")
            markdown_content.append(
                f"**Chapter {chapter_num} - Section {section_num}**\n\n"
            )
            markdown_content.append("---\n\n")

            # Process elements
            para_count = 0
            table_count = 0

            for elem_type, elem in section_elements:
                if elem_type == "paragraph":
                    if should_skip_toc_paragraph(elem):
                        continue

                    para_md, chapter_image_counter = extract_paragraph_markdown(
                        elem, chapter_image_counter, pictures_dir, doc, chapter_num
                    )

                    if para_md.strip():
                        markdown_content.append(para_md)
                        markdown_content.append("\n")
                        para_count += 1

                elif elem_type == "table":
                    table_md = extract_table_markdown(elem)
                    if table_md.strip():
                        markdown_content.append(table_md)
                        markdown_content.append("\n")
                        table_count += 1

            # Save section markdown file
            section_filename = f"section_{section_num.replace('.', '_')}.md"
            section_filepath = os.path.join(chapter_dir, section_filename)

            with open(section_filepath, "w", encoding="utf-8") as f:
                f.write("".join(markdown_content))

            section_files.append(
                {
                    "number": section_num,
                    "title": section_title,
                    "filename": section_filename,
                    "paragraphs": para_count,
                    "tables": table_count,
                }
            )

            total_sections += 1
            total_paragraphs += para_count
            total_tables += table_count

        total_images += chapter_image_counter - 1

        # Count WMF conversions in this chapter
        for image_file in Path(pictures_dir).glob("*.wmf.backup"):
            total_wmf_converted += 1

        # Create chapter index file
        chapter_index_md = []
        chapter_index_md.append(f"# {chapter_title}\n\n")
        chapter_index_md.append(f"**Chapter {chapter_num}**\n\n")
        chapter_index_md.append("## Sections\n\n")

        for section_file in section_files:
            chapter_index_md.append(
                f"- [{section_file['number']} {section_file['title']}](./{section_file['filename']})\n"
            )

        chapter_index_md.append(f"\n## Statistics\n\n")
        chapter_index_md.append(f"- **Sections**: {len(section_files)}\n")
        chapter_index_md.append(
            f"- **Paragraphs**: {sum(s['paragraphs'] for s in section_files)}\n"
        )
        chapter_index_md.append(
            f"- **Tables**: {sum(s['tables'] for s in section_files)}\n"
        )
        chapter_index_md.append(f"- **Images**: {chapter_image_counter - 1}\n")

        with open(os.path.join(chapter_dir, "README.md"), "w", encoding="utf-8") as f:
            f.write("".join(chapter_index_md))

        # Add to main index
        chapter_index.append(
            {
                "number": chapter_num,
                "title": chapter_title,
                "sections": len(section_files),
                "directory": f"chapter_{chapter_num:02d}",
            }
        )

        print_success(
            f"    Chapter {chapter_num}: {len(section_files)} sections, "
            f"{sum(s['paragraphs'] for s in section_files)} paragraphs, "
            f"{sum(s['tables'] for s in section_files)} tables"
        )

    # Create main index
    print_step(5, "Creating main index")

    main_index_md = []
    main_index_md.append("# Book Contents - Markdown Export\n\n")
    main_index_md.append(f"**Source**: {input_file}\n\n")
    main_index_md.append("---\n\n")
    main_index_md.append("## Chapters\n\n")

    for chapter in chapter_index:
        main_index_md.append(
            f"{chapter['number']}. [{chapter['title']}](./{chapter['directory']}/README.md) "
            f"({chapter['sections']} sections)\n"
        )

    main_index_md.append(f"\n## Overall Statistics\n\n")
    main_index_md.append(f"- **Total Chapters**: {len(chapters)}\n")
    main_index_md.append(f"- **Total Sections**: {total_sections}\n")
    main_index_md.append(f"- **Total Paragraphs**: {total_paragraphs}\n")
    main_index_md.append(f"- **Total Tables**: {total_tables}\n")
    main_index_md.append(f"- **Total Images**: {total_images}\n")

    main_index_md.append(f"\n## About This Export\n\n")
    main_index_md.append(
        "This markdown export was generated by `split_to_md_chapters.py`, "
    )
    main_index_md.append(
        "a reference implementation that demonstrates the document conversion logic.\n\n"
    )
    main_index_md.append("### Features\n\n")
    main_index_md.append(
        "- **Formatted Text**: Bold, italic, and underline formatting preserved\n"
    )
    main_index_md.append("- **Tables**: Converted to Markdown table format\n")
    main_index_md.append("- **Images**: Extracted and referenced with proper paths\n")
    main_index_md.append("- **Structure**: Organized by chapters and sections\n")
    main_index_md.append(
        "- **Navigation**: Each chapter has an index for easy browsing\n\n"
    )
    main_index_md.append("### Viewing\n\n")
    main_index_md.append("You can view these files in:\n")
    main_index_md.append("- Any Markdown viewer (VS Code, Typora, MacDown, etc.)\n")
    main_index_md.append("- GitHub (renders Markdown automatically)\n")
    main_index_md.append("- Static site generators (Jekyll, Hugo, MkDocs, etc.)\n")

    with open(os.path.join(output_dir, "README.md"), "w", encoding="utf-8") as f:
        f.write("".join(main_index_md))

    print_success("Created main index: README.md")

    # Summary
    print_header("CONVERSION COMPLETE")
    print_success(f"Generated {len(chapters)} chapters")
    print_success(f"Total sections: {total_sections}")
    print_success(f"Total paragraphs: {total_paragraphs}")
    print_success(f"Total tables: {total_tables}")
    print_success(f"Total images: {total_images}")
    if total_wmf_converted > 0:
        print_success(f"Converted {total_wmf_converted} WMF images to PNG")
    print_info(f"\nMarkdown files saved to: {output_dir}/")
    print_info("Open README.md to start browsing")


# ============================================================================
# Main Entry Point
# ============================================================================


def main():
    """Main entry point."""

    if not os.path.exists(INPUT_DOCX):
        print_error(f"Input file not found: {INPUT_DOCX}")
        print_info("Please ensure the Word document is in the current directory")
        sys.exit(1)

    try:
        convert_book_to_markdown(INPUT_DOCX, MARKDOWN_DIR)

        print("\n" + "=" * 70)
        print(f"{Colors.GREEN}{Colors.BOLD}SUCCESS!{Colors.END}")
        print("=" * 70)
        print(f"\n{Colors.CYAN}Next steps:{Colors.END}")
        print(f"  1. Open {MARKDOWN_DIR}/README.md")
        print(f"  2. Browse chapters in any Markdown viewer")
        print(f"  3. Compare with chapter-viewer output\n")

    except Exception as e:
        print_error(f"Error during conversion: {e}")
        import traceback

        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
