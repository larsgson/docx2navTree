#!/usr/bin/env python3
"""
Script to split a book (DOCX) into separate chapter files.
Chapters are identified by bold headings starting with N.0 format (e.g., 1.0, 2.0, etc.)
where each chapter starts on a new page.

This script preserves:
- Text formatting (bold, italic, underline, fonts, colors)
- Pictures/images in their correct position within text flow (with actual image data)
- Tables in their correct position within text flow
- Footnotes associated with each chapter
"""

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches
import re
import os
from io import BytesIO
from copy import deepcopy


def is_chapter_heading(paragraph):
    """
    Check if a paragraph is a chapter heading.
    Chapter headings are bold text starting with N.0 format.
    """
    text = paragraph.text.strip()
    if not re.match(r'^\d+\.0\s', text):
        return False
    for run in paragraph.runs:
        if run.bold and run.text.strip():
            return True
    return False


def extract_chapter_number(paragraph):
    """Extract chapter number from heading (e.g., "1.0 Chapter Title" -> "1")"""
    text = paragraph.text.strip()
    match = re.match(r'^(\d+)\.0\s', text)
    if match:
        return match.group(1)
    return None


def get_footnote_references_from_paragraph(paragraph):
    """Get all footnote reference IDs from a paragraph."""
    footnote_refs = []
    try:
        refs = paragraph._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnoteReference')
        for footnote_ref in refs:
            footnote_id = footnote_ref.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
            if footnote_id:
                footnote_refs.append((footnote_id, footnote_ref))
        
        refs = paragraph._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}endnoteReference')
        for endnote_ref in refs:
            endnote_id = endnote_ref.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
            if endnote_id:
                footnote_refs.append((endnote_id, endnote_ref))
    except:
        pass
    return footnote_refs


def get_footnote_references_from_table(table):
    """Get all footnote references from a table."""
    footnote_refs = []
    try:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    refs = get_footnote_references_from_paragraph(para)
                    footnote_refs.extend(refs)
    except:
        pass
    return footnote_refs


def copy_paragraph_with_images(source_para, target_doc, source_doc):
    """Copy a paragraph preserving all formatting, images, and footnotes."""
    new_para = target_doc.add_paragraph()
    new_para.alignment = source_para.alignment
    new_para.style = source_para.style
    
    if source_para.paragraph_format:
        try:
            new_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
            new_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
            new_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
            new_para.paragraph_format.space_before = source_para.paragraph_format.space_before
            new_para.paragraph_format.space_after = source_para.paragraph_format.space_after
            new_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
        except:
            pass
    
    for run in source_para.runs:
        has_image = False
        
        for drawing in run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
            has_image = True
            try:
                blip = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                if blip is not None:
                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    
                    if embed_id:
                        image_part = source_doc.part.related_parts[embed_id]
                        image_bytes = image_part.blob
                        
                        extent = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                        if extent is not None:
                            width = int(extent.get('cx', 914400))
                            height = int(extent.get('cy', 914400))
                            width_inches = width / 914400.0
                            height_inches = height / 914400.0
                            
                            max_width = 6.0
                            max_height = 9.0
                            
                            if width_inches > max_width:
                                ratio = max_width / width_inches
                                width_inches = max_width
                                height_inches *= ratio
                            
                            if height_inches > max_height:
                                ratio = max_height / height_inches
                                height_inches = max_height
                                width_inches *= ratio
                        else:
                            width_inches = 3.0
                            height_inches = 3.0
                        
                        new_run = new_para.add_run()
                        new_run.add_picture(BytesIO(image_bytes), width=Inches(width_inches), height=Inches(height_inches))
            except:
                new_run = new_para.add_run("[Image]")
        
        if not has_image:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name
            try:
                if run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
            except:
                pass
            
            try:
                for footnote_ref in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnoteReference'):
                    new_footnote_ref = deepcopy(footnote_ref)
                    new_run._element.append(new_footnote_ref)
            except:
                pass
    
    return new_para


def copy_table(source_table, target_doc):
    """Copy a table preserving content and basic structure."""
    new_table = target_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    
    try:
        if source_table.style:
            new_table.style = source_table.style
    except:
        pass
    
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            new_cell = new_table.rows[i].cells[j]
            
            if new_cell.paragraphs:
                new_cell.paragraphs[0].clear()
            
            for k, para in enumerate(cell.paragraphs):
                if k == 0:
                    new_para = new_cell.paragraphs[0]
                else:
                    new_para = new_cell.add_paragraph()
                
                new_para.alignment = para.alignment
                try:
                    new_para.style = para.style
                except:
                    pass
                
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    if run.font.size:
                        new_run.font.size = run.font.size
                    if run.font.name:
                        new_run.font.name = run.font.name
                    
                    try:
                        for footnote_ref in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnoteReference'):
                            new_footnote_ref = deepcopy(footnote_ref)
                            new_run._element.append(new_footnote_ref)
                    except:
                        pass
    
    return new_table


def get_document_elements(doc):
    """Get all document elements (paragraphs and tables) in order."""
    elements = []
    for element in doc.element.body:
        if isinstance(element, CT_P):
            para = Paragraph(element, doc)
            elements.append(('paragraph', para))
        elif isinstance(element, CT_Tbl):
            table = Table(element, doc)
            elements.append(('table', table))
    return elements


def collect_chapter_footnotes(chapter_elements):
    """Collect all footnote IDs referenced in the chapter."""
    footnote_map = {}
    
    for elem_type, elem in chapter_elements:
        if elem_type == 'paragraph':
            refs = get_footnote_references_from_paragraph(elem)
            for footnote_id, footnote_elem in refs:
                if footnote_id not in footnote_map:
                    footnote_map[footnote_id] = []
                footnote_map[footnote_id].append(('paragraph', elem))
        elif elem_type == 'table':
            refs = get_footnote_references_from_table(elem)
            for footnote_id, footnote_elem in refs:
                if footnote_id not in footnote_map:
                    footnote_map[footnote_id] = []
                footnote_map[footnote_id].append(('table', elem))
    
    return footnote_map


def copy_footnotes_to_document(footnote_ids, source_doc, target_doc):
    """Copy footnotes from source to target document."""
    if not footnote_ids:
        return
    
    try:
        footnotes_part = None
        for rel in source_doc.part.rels.values():
            if 'footnotes' in rel.target_ref.lower():
                footnotes_part = rel.target_part
                break
        
        if footnotes_part is None:
            return
        
        from lxml import etree
        if hasattr(footnotes_part, 'element'):
            footnotes_element = footnotes_part.element
        elif hasattr(footnotes_part, '_element'):
            footnotes_element = footnotes_part._element
        else:
            footnotes_element = etree.fromstring(footnotes_part.blob)
        
        new_footnotes_element = deepcopy(footnotes_element)
        
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.part import XmlPart
        from docx.opc.packuri import PackURI
        
        partname = '/word/footnotes.xml'
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
        
        pkg = target_doc.part.package
        part_uri = PackURI(partname)
        target_footnotes_part = XmlPart(part_uri, content_type, new_footnotes_element, pkg)
        target_doc.part.relate_to(target_footnotes_part, RT.FOOTNOTES)
    
    except:
        pass


def read_toc_from_document(doc):
    """Read table of contents from the document for reference."""
    toc_entries = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if re.match(r'^\d+\.0\s', text):
            toc_entries.append(text)
    return toc_entries


def split_book_into_chapters(input_file, output_dir="chapters"):
    """
    Split the book into separate chapter files, preserving paragraphs, tables, images, and footnotes.
    """
    os.makedirs(output_dir, exist_ok=True)
    
    print(f"Loading document: {input_file}")
    doc = Document(input_file)
    
    elements = get_document_elements(doc)
    print(f"Total elements in document: {len(elements)}")
    
    chapters = []
    current_chapter = None
    current_chapter_title = None
    current_chapter_elements = []
    
    print("Processing document elements...")
    
    for elem_type, elem in elements:
        if elem_type == 'paragraph' and is_chapter_heading(elem):
            chapter_num = extract_chapter_number(elem)
            
            # Skip TOC entries (short chapters with < 50 elements)
            if current_chapter == chapter_num and len(current_chapter_elements) < 50:
                current_chapter = None
                current_chapter_title = None
                current_chapter_elements = []
                continue
            
            # Save previous chapter if it has substantial content
            if current_chapter is not None and len(current_chapter_elements) > 10:
                chapters.append({
                    'number': current_chapter,
                    'title': current_chapter_title,
                    'elements': current_chapter_elements
                })
                print(f"Found Chapter {current_chapter}: {current_chapter_title[:60]}... ({len(current_chapter_elements)} elements)")
            
            current_chapter = chapter_num
            current_chapter_title = elem.text.strip()
            current_chapter_elements = [(elem_type, elem)]
        else:
            if current_chapter is not None:
                current_chapter_elements.append((elem_type, elem))
    
    # Save last chapter
    if current_chapter is not None and len(current_chapter_elements) > 10:
        chapters.append({
            'number': current_chapter,
            'title': current_chapter_title,
            'elements': current_chapter_elements
        })
        print(f"Found Chapter {current_chapter}: {current_chapter_title[:60]}... ({len(current_chapter_elements)} elements)")
    
    print(f"\nTotal chapters found: {len(chapters)}")
    
    # Sort chapters by chapter number
    chapters.sort(key=lambda x: int(x['number']))
    
    # Create separate files for each chapter
    print("\nCreating chapter files...")
    for chapter in chapters:
        print(f"\nProcessing Chapter {chapter['number']}: {chapter['title'][:60]}...")
        chapter_num = chapter['number'].zfill(2)
        filename = f"chapter_{chapter_num}.docx"
        filepath = os.path.join(output_dir, filename)
        
        # Collect footnotes for this chapter
        footnote_map = collect_chapter_footnotes(chapter['elements'])
        
        # Create new document for this chapter
        chapter_doc = Document()
        
        # Copy elements to new document in order
        para_count = 0
        table_count = 0
        image_count = 0
        
        for elem_type, elem in chapter['elements']:
            if elem_type == 'paragraph':
                for run in elem.runs:
                    if run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
                        image_count += 1
                
                copy_paragraph_with_images(elem, chapter_doc, doc)
                para_count += 1
            elif elem_type == 'table':
                copy_table(elem, chapter_doc)
                table_count += 1
        
        # Copy footnotes if any
        if footnote_map:
            copy_footnotes_to_document(list(footnote_map.keys()), doc, chapter_doc)
        
        # Save chapter document
        chapter_doc.save(filepath)
        print(f"  Saved: {filename} ({para_count} paragraphs, {table_count} tables, {image_count} images, {len(footnote_map)} footnotes)")
    
    print(f"\nDone! {len(chapters)} chapters saved to '{output_dir}/' directory")
    return chapters


if __name__ == "__main__":
    BOOK_FILE = "English HAH Word Apr 6 2024.docx"
    OUTPUT_DIR = "chapters"
    
    print("=" * 60)
    print("Book Chapter Splitter")
    print("=" * 60)
    print()
    
    try:
        chapters = split_book_into_chapters(BOOK_FILE, OUTPUT_DIR)
        print("\n" + "=" * 60)
        print("SUCCESS!")
        print("=" * 60)
    except Exception as e:
        print(f"\nError: {e}")
        import traceback
        traceback.print_exc()